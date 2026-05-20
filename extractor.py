#!/usr/bin/env python3
"""
Document Extraction Pipeline - Single Faculty
===============================================
Extracts content from all files inside a given faculty folder
without deleting or modifying original files.

Usage:
  python pipeline_single.py --input "/path/to/faculty_folder" --output "extracted"

Filter logic (step 2):
  - A file is KEPT only if it is a schedule (emploi de temps / emploi d'examens):
      → contains "emploi" or "horaire" in its text  OR
      → contains at least one table  AND  contains day/time-like content
  - A file is DROPPED if:
      → it contains any course keywords (cours, td, serie, exam, etc.)
      → OR it matches a noise keyword AND does not qualify as a schedule
  - A file is also DROPPED if its extracted text is shorter than MIN_TEXT_CHARS
    (step 3: no meaningful content)
"""

import sys, json, hashlib, logging, re, shutil, tempfile, zipfile, argparse
import unicodedata
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Tuple, Optional

# ── optional heavy deps ───────────────────────────────────────────────────────
try:
    import fitz
    FITZ_OK = True
except ImportError:
    print("ERROR: PyMuPDF not installed.  Run: pip install pymupdf")
    sys.exit(1)

try:
    import pdfplumber
    PDFPLUMBER_OK = True
except ImportError:
    print("WARNING: pdfplumber not installed – PyMuPDF fallback for tables.")
    PDFPLUMBER_OK = False

try:
    from docx import Document as DocxDocument
    DOCX_OK = True
except ImportError:
    print("WARNING: python-docx not installed – DOCX files skipped.")
    DOCX_OK = False

try:
    import pytesseract
    from PIL import Image, ImageEnhance, ImageFilter
    OCR_OK = True
except ImportError:
    print("WARNING: pytesseract/Pillow not installed – image OCR disabled.")
    OCR_OK = False

try:
    import openpyxl
    OPENPYXL_OK = True
except ImportError:
    print("WARNING: openpyxl not installed – .xlsx files skipped.")
    OPENPYXL_OK = False

try:
    import xlrd
    XLRD_OK = True
except ImportError:
    XLRD_OK = False

try:
    import rarfile
    RAR_OK = True
except ImportError:
    print("WARNING: rarfile not installed – .rar files skipped.")
    RAR_OK = False

# ── logging ───────────────────────────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
)
log = logging.getLogger(__name__)

# ── extension maps ────────────────────────────────────────────────────────────
PDF_EXT     = {".pdf"}
DOCX_EXT    = {".docx", ".doc"}
IMAGE_EXT   = {".png", ".jpg", ".jpeg", ".tiff", ".tif", ".bmp", ".webp"}
XLSX_EXT    = {".xlsx", ".xls"}
ARCHIVE_EXT = {".zip", ".rar"}

EXTRACTABLE_EXT = PDF_EXT | DOCX_EXT | IMAGE_EXT | XLSX_EXT
ALL_EXT         = EXTRACTABLE_EXT | ARCHIVE_EXT

MAX_ARCHIVE_DEPTH = 5

# ── Arabic text marker ────────────────────────────────────────────────────────
ARABIC_RE = re.compile(
    r"[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF\uFB50-\uFDFF\uFE70-\uFEFF]"
)

# ── Minimum characters to consider a file "meaningful" (step 3) ──────────────
MIN_TEXT_CHARS = 50

def normalize_text(text: str) -> str:

    if not text:
        return ""

    text = text.lower()

    text = unicodedata.normalize("NFKD", text)
    text = "".join(
        c for c in text
        if not unicodedata.combining(c)
    )

    text = re.sub(r"[_\-.\\/]+", " ", text)

    text = re.sub(r"([a-z])([A-Z])", r"\1 \2", text)

    text = re.sub(r"(\d)([a-z])", r"\1 \2", text)
    text = re.sub(r"([a-z])(\d)", r"\1 \2", text)

    text = re.sub(r"\s+", " ", text)

    return text.strip()


def normalize_filename(path: Path) -> str:
    return normalize_text(path.stem)

# ──────────────────────────────────────────────────────────────────────────────
# KEYWORDS
# ──────────────────────────────────────────────────────────────────────────────

COURSE_KEYWORDS = {

    # courses/books
    "cours",
    "course",
    "lecture",
    "lesson",
    "chapter",
    "chapitre",
    "module",
    "support",
    "book",
    "ebook",
    "manuel",
    "handbook",
    "notes",
    "slides",
    "presentation",
    "resume",
    "summary",

    # exercises
    "td",
    "tp",
    "tutorial",
    "exercise",
    "exercises",
    "exercice",
    "exercices",
    "serie",
    "series",
    "worksheet",
    "problem",
    "lab",
    "practical",

    # exams
    "exam",
    "test",
    "quiz",
    "controle",
    "interrogation",
    "midterm",
    "final",
    "rattrapage",
    "ratrappage",
    "sujet",
    "corrige",
    "corrigé",
    "solution",
    "annale",
    "annales",

    # assignments
    "devoir",
    "assignment",
    "homework",
    "project",
    "projet",

    # arabic
    "محاضرات",
    "دروس",
    "تمارين",
    "اعمال",
}

COMMON_SUBSTRINGS = [
    "cour",
    "chap",
    "td",
    "tp",
    "exo",
    "corr",
    "rattr",
    "annal",
    "interro",
    "devoir",
    "sujet",
    "serie",
    "book",
    "exam",
    "quiz",
]

# ──────────────────────────────────────────────────────────────────────────────
# SCHEDULE DETECTION
# ──────────────────────────────────────────────────────────────────────────────

SCHEDULE_WORDS = re.compile(
    r"(emploi|horaire|timetable|planning|schedule|"
    r"emploi du temps|emploi des examens|creneau|seance)",
    re.IGNORECASE
)

DAY_WORDS = re.compile(
    r"(lundi|mardi|mercredi|jeudi|vendredi|samedi|dimanche|"
    r"monday|tuesday|wednesday|thursday|friday|saturday|sunday|"
    r"الاحد|الاثنين|الثلاثاء|الاربعاء|الخميس|الجمعة|السبت)",
    re.IGNORECASE
)

TIME_PATTERN = re.compile(
    r"\b\d{1,2}[:h]\d{0,2}\b",
    re.IGNORECASE
)

# ──────────────────────────────────────────────────────────────────────────────
# HELPERS
# ──────────────────────────────────────────────────────────────────────────────

def contains_keyword(text: str, keywords: set) -> bool:

    if not text:
        return False

    text = normalize_text(text)

    for kw in keywords:
        if kw in text:
            return True

    return False


def contains_partial_keyword(text: str) -> bool:

    if not text:
        return False

    text = normalize_text(text)

    for sub in COMMON_SUBSTRINGS:
        if sub in text:
            return True

    return False


def is_schedule(
    text: str,
    tables: list,
    file_name: str = "",
) -> bool:

    combined = normalize_text(
        f"{file_name} {text}"
    )

    has_schedule_word = bool(
        SCHEDULE_WORDS.search(combined)
    )

    has_days = bool(
        DAY_WORDS.search(combined)
    )

    has_times = bool(
        TIME_PATTERN.search(combined)
    )

    has_tables = bool(tables)

    score = 0

    if has_schedule_word:
        score += 4

    if has_days:
        score += 2

    if has_times:
        score += 2

    if has_tables:
        score += 1

    return score >= 4


def has_course_keywords(text: str, file_name: str = "") -> Tuple[bool, Optional[str]]:
    """
    Check if the text or filename contains any course keywords.
    Returns (True, keyword_found) or (False, None).
    """
    combined = normalize_text(f"{file_name} {text}")
    
    # Check for exact keywords
    for kw in COURSE_KEYWORDS:
        if kw in combined:
            return True, kw
    
    # Check for common substrings
    for sub in COMMON_SUBSTRINGS:
        if sub in combined:
            return True, sub
    
    return False, None


def classify_document(
    text: str,
    file_name: str = "",
    tables: list = None,
) -> Tuple[str, int]:

    if tables is None:
        tables = []

    # First, check for course keywords (these will be dropped)
    has_course, keyword = has_course_keywords(text, file_name)
    if has_course:
        return "course_material", 10

    normalized_text = normalize_text(text)
    normalized_name = normalize_text(file_name)

    combined = f"{normalized_name} {normalized_text}"

    if is_schedule(text, tables, file_name):
        return "schedule", 100

    if not is_meaningful(text):
        return "too_short", 0

    return "other", 0


def is_meaningful(text: str) -> bool:

    if not text:
        return False

    return len(normalize_text(text)) >= MIN_TEXT_CHARS


def should_keep(
    text: str,
    tables: list,
    file_name: str = "",
) -> Tuple[bool, str]:

    doc_type, score = classify_document(
        text=text,
        file_name=file_name,
        tables=tables,
    )

    # ONLY keep schedule files
    if doc_type == "schedule":
        return True, f"schedule detected (score={score})"
    
    # Drop everything else (course materials, short texts, other documents)
    if doc_type == "course_material":
        return False, f"course material detected - contains course keyword"
    
    if doc_type == "too_short":
        return False, (
            f"text too short "
            f"({len(normalize_text(text))} chars < {MIN_TEXT_CHARS})"
        )
    
    return False, f"not a schedule (type={doc_type})"

# ╔══════════════════════════════════════════════════════════════════════════╗
# ║  PHASE 1 — ARCHIVE UNPACKER & FILE SORTER (NO DELETION)                  ║
# ╚══════════════════════════════════════════════════════════════════════════╝

def folder_for_ext(ext: str, input_dir: Path) -> Optional[Path]:
    """Return the target subfolder for a given extension."""
    if ext in PDF_EXT:   return input_dir / "pdfs"
    if ext in DOCX_EXT:  return input_dir / "docs"
    if ext in IMAGE_EXT: return input_dir / "images"
    if ext in XLSX_EXT:  return input_dir / "docs"   # xlsx/xls live alongside docs
    return None


def _unpack_zip(archive_path: Path, dest_dir: Path) -> bool:
    try:
        with zipfile.ZipFile(archive_path, "r") as zf:
            for member in zf.infolist():
                mp = Path(member.filename)
                if mp.is_absolute() or ".." in mp.parts:
                    log.warning("Skipping unsafe ZIP entry: %s", member.filename)
                    continue
                zf.extract(member, dest_dir)
        return True
    except Exception as exc:
        log.error("ZIP unpack failed %s: %s", archive_path.name, exc)
        return False


def _unpack_rar(archive_path: Path, dest_dir: Path) -> bool:
    if not RAR_OK:
        log.warning("rarfile not installed – skipping %s", archive_path.name)
        return False
    try:
        with rarfile.RarFile(str(archive_path)) as rf:
            rf.extractall(str(dest_dir))
        return True
    except Exception as exc:
        log.error("RAR unpack failed %s: %s", archive_path.name, exc)
        return False


def copy_and_sort(src: Path, input_dir: Path) -> Optional[Path]:
    """
    COPY *src* into the correct subfolder of input_dir (keeps original intact).
    Returns the new path, or None if unsupported / already in target folder.
    """
    ext = src.suffix.lower()
    dest_folder = folder_for_ext(ext, input_dir)
    if dest_folder is None:
        log.debug("No target folder for %s – skipping", src.name)
        return None

    dest_folder.mkdir(parents=True, exist_ok=True)
    dest = dest_folder / src.name

    # Avoid overwriting: append size suffix if name collides
    if dest.exists():
        dest = dest_folder / f"{src.stem}_{src.stat().st_size}{src.suffix}"

    shutil.copy2(str(src), dest)
    log.info("    → copied %s  →  %s/", src.name, dest_folder.name)
    return dest


def _sort_root_files(input_dir: Path) -> None:
    """
    Copy any extractable files sitting directly in input_dir into the
    appropriate subfolders (pdfs / docs / images).
    Archives are also copied to docs/ so Phase 1 can unpack them.
    """
    for f in list(input_dir.iterdir()):
        if not f.is_file():
            continue
        ext = f.suffix.lower()
        if ext in EXTRACTABLE_EXT:
            copy_and_sort(f, input_dir)
        elif ext in ARCHIVE_EXT:
            dest_folder = input_dir / "docs"
            dest_folder.mkdir(parents=True, exist_ok=True)
            dest = dest_folder / f.name
            if not dest.exists():
                shutil.copy2(str(f), dest)
                log.info("    → root archive copied to docs/: %s", f.name)


def unpack_archives_keep(input_dir: Path, depth: int = 0) -> List[Path]:
    """
    Recursively unpack archives found in input_dir/docs/.
    Files are COPIED to the correct subfolders (pdfs, docs, images).
    Archives themselves are NEVER deleted.
    Returns a list of all files that were extracted (copied).
    """
    if depth >= MAX_ARCHIVE_DEPTH:
        log.warning("Max archive depth reached, stopping recursion.")
        return []

    doc_folder = input_dir / "docs"
    if not doc_folder.exists():
        return []

    archives = [
        f for f in doc_folder.iterdir()
        if f.is_file() and f.suffix.lower() in ARCHIVE_EXT
    ]

    if not archives:
        log.info("  No archives found in %s", doc_folder)
        return []

    extracted_files: List[Path] = []
    tmp_base = Path(tempfile.mkdtemp(prefix="pipeline_"))

    try:
        for archive in sorted(archives):
            log.info("[UNPACK] %s", archive.name)
            tmp = tmp_base / archive.stem
            tmp.mkdir(parents=True, exist_ok=True)

            ext = archive.suffix.lower()
            ok = _unpack_zip(archive, tmp) if ext == ".zip" else _unpack_rar(archive, tmp)
            if not ok:
                continue

            log.info("  ✓ Unpacked (original archive kept): %s", archive.name)

            nested_archives_found = False
            for inner in sorted(tmp.rglob("*")):
                if not inner.is_file():
                    continue
                inner_ext = inner.suffix.lower()

                if inner_ext in ARCHIVE_EXT:
                    dest_folder = input_dir / "docs"
                    dest_folder.mkdir(parents=True, exist_ok=True)
                    dest = dest_folder / inner.name
                    if dest.exists():
                        dest = dest_folder / f"{inner.stem}_{inner.stat().st_size}{inner.suffix}"
                    shutil.copy2(str(inner), dest)
                    log.info("    → nested archive copied: %s", inner.name)
                    nested_archives_found = True

                elif inner_ext in EXTRACTABLE_EXT:
                    new_path = copy_and_sort(inner, input_dir)
                    if new_path:
                        extracted_files.append(new_path)
                else:
                    log.debug("    – skipping unsupported inner file: %s", inner.name)

            if nested_archives_found:
                extracted_files.extend(unpack_archives_keep(input_dir, depth + 1))

    finally:
        shutil.rmtree(tmp_base, ignore_errors=True)

    return extracted_files


# ╔══════════════════════════════════════════════════════════════════════════╗
# ║  PHASE 2 — CONTENT EXTRACTION                                            ║
# ╚══════════════════════════════════════════════════════════════════════════╝

def _fix_arabic(text: str) -> str:
    if not text or not ARABIC_RE.search(text):
        return text
    return "\n".join(
        f"\u200F{ln}\u200F" if ARABIC_RE.search(ln) else ln
        for ln in text.split("\n")
    )

def _file_id(path: Path) -> str:
    return hashlib.md5(str(path).encode()).hexdigest()[:8]

def _cell_value(cell) -> str:
    return "" if cell.value is None else str(cell.value).strip()


# ── PDF tables ────────────────────────────────────────────────────────────────

class PDFPlumberTableExtractor:
    @staticmethod
    def _normalize_row(row, n):
        return (list(row) + [None] * n)[:n]

    @staticmethod
    def _fix_colspan(row):
        out, last = [], None
        for cell in row:
            if cell is None or str(cell).strip() == "":
                out.append(last)
            else:
                out.append(cell); last = cell
        return out

    @staticmethod
    def _fix_rowspan(rows):
        for i in range(1, len(rows)):
            for j in range(len(rows[i])):
                if rows[i][j] is None or str(rows[i][j]).strip() == "":
                    rows[i][j] = rows[i-1][j]
        return rows

    @staticmethod
    def _is_valid(t):
        return t and len(t) >= 2 and any(
            c is not None and str(c).strip() for r in t for c in r)

    @staticmethod
    def _clean_headers(h):
        return [str(v).strip() if v and str(v).strip() else f"column_{i}"
                for i, v in enumerate(h)]

    def extract(self, pdf_path: Path) -> List[Dict]:
        if not PDFPLUMBER_OK: return []
        results = []
        try:
            with pdfplumber.open(pdf_path) as pdf:
                for p_idx, page in enumerate(pdf.pages):
                    for t_idx, table in enumerate(page.extract_tables() or []):
                        if not self._is_valid(table): continue
                        headers = self._clean_headers(table[0])
                        n = len(headers)
                        rows = [self._fix_colspan(self._normalize_row(r, n))
                                for r in table[1:]]
                        rows = self._fix_rowspan(rows)
                        clean = [
                            {headers[i]: (r[i] or "") for i in range(n)}
                            for r in rows if any(v is not None and str(v).strip() for v in r)
                        ]
                        if clean:
                            results.append({"type":"table","page":p_idx+1,
                                "table_index":t_idx,"headers":headers,
                                "rows":clean,"extraction_method":"pdfplumber"})
        except Exception as exc:
            log.debug("pdfplumber error %s: %s", pdf_path.name, exc)
        return results


class PyMuPDFTableExtractor:
    def extract(self, pdf_path: Path) -> List[Dict]:
        results = []
        try:
            doc = fitz.open(pdf_path)
            for p_idx in range(len(doc)):
                page = doc[p_idx]
                try: found = page.find_tables()
                except AttributeError: continue
                if not found or not found.tables: continue
                for t_idx, table in enumerate(found.tables):
                    data = []
                    for r in range(table.row_count):
                        data.append([
                            (table.cell((r,c)).text.strip() if table.cell((r,c)) else "")
                            for c in range(table.col_count)])
                    if not data: continue
                    headers = data[0]
                    clean = [{headers[i]:(row[i] if i<len(row) else "")
                              for i in range(len(headers))} for row in data[1:]]
                    results.append({"type":"table","page":p_idx+1,
                        "table_index":t_idx,"headers":headers,
                        "rows":clean,"extraction_method":"pymupdf"})
            doc.close()
        except Exception as exc:
            log.debug("pymupdf table error %s: %s", pdf_path.name, exc)
        return results


# ── XLSX ──────────────────────────────────────────────────────────────────────

def extract_xlsx(xlsx_path: Path) -> Tuple[str, List[Dict]]:
    ext = xlsx_path.suffix.lower()
    tables: List[Dict] = []; text_parts: List[str] = []

    if ext == ".xlsx" and OPENPYXL_OK:
        try:
            wb = openpyxl.load_workbook(xlsx_path, data_only=True)
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                all_rows = list(ws.iter_rows())
                if not all_rows: continue
                raw_headers = [_cell_value(c) for c in all_rows[0]]
                headers = [h if h else f"column_{i}" for i,h in enumerate(raw_headers)]
                n = len(headers); data_rows=[]; tsv_lines=["\t".join(headers)]
                for row in all_rows[1:]:
                    cells = [_cell_value(c) for c in row]
                    if not any(cells): continue
                    cells = (cells+[""]*n)[:n]
                    data_rows.append({headers[i]:cells[i] for i in range(n)})
                    tsv_lines.append("\t".join(cells))
                if data_rows:
                    tables.append({"type":"table","sheet":sheet_name,
                        "table_index":len(tables),"headers":headers,
                        "rows":data_rows,"extraction_method":"openpyxl"})
                    text_parts.append(f"[SHEET: {sheet_name}]\n"+"\n".join(tsv_lines))
            wb.close()
        except Exception as exc:
            log.error("openpyxl error %s: %s", xlsx_path.name, exc)

    elif ext == ".xls" and XLRD_OK:
        try:
            wb = xlrd.open_workbook(str(xlsx_path))
            for sheet_name in wb.sheet_names():
                ws = wb.sheet_by_name(sheet_name)
                if ws.nrows == 0: continue
                raw_headers = [str(ws.cell_value(0,c)).strip() for c in range(ws.ncols)]
                headers = [h if h else f"column_{i}" for i,h in enumerate(raw_headers)]
                n = len(headers); data_rows=[]; tsv_lines=["\t".join(headers)]
                for r in range(1, ws.nrows):
                    cells = [str(ws.cell_value(r,c)).strip() for c in range(ws.ncols)]
                    if not any(cells): continue
                    cells = (cells+[""]*n)[:n]
                    data_rows.append({headers[i]:cells[i] for i in range(n)})
                    tsv_lines.append("\t".join(cells))
                if data_rows:
                    tables.append({"type":"table","sheet":sheet_name,
                        "table_index":len(tables),"headers":headers,
                        "rows":data_rows,"extraction_method":"xlrd"})
                    text_parts.append(f"[SHEET: {sheet_name}]\n"+"\n".join(tsv_lines))
        except Exception as exc:
            log.error("xlrd error %s: %s", xlsx_path.name, exc)

    return _fix_arabic("\n\n".join(text_parts)), tables


# ── PDF ───────────────────────────────────────────────────────────────────────

def extract_pdf(pdf_path: Path) -> Tuple[str, List[Dict], List[Dict]]:
    tables = PDFPlumberTableExtractor().extract(pdf_path)
    if not tables:
        tables = PyMuPDFTableExtractor().extract(pdf_path)
    pages=[]; full_parts=[]
    try:
        doc = fitz.open(pdf_path)
        for page_num, page in enumerate(doc, start=1):
            page_text = page.get_text() or ""
            pages.append({"page":page_num,"text":page_text.strip()})
            full_parts.append(page_text)
        doc.close()
    except Exception as exc:
        log.error("PyMuPDF text error %s: %s", pdf_path.name, exc)
    return _fix_arabic("\n".join(full_parts)), pages, tables


# ── DOCX ──────────────────────────────────────────────────────────────────────

def extract_docx(docx_path: Path) -> Tuple[str, List[Dict], List[Dict]]:
    if not DOCX_OK:
        log.warning("python-docx not available, skipping %s", docx_path.name)
        return "", [], []
    tables=[]; text_parts=[]; embedded_images=[]
    try:
        doc = DocxDocument(docx_path)
        for para in doc.paragraphs:
            if para.text.strip(): text_parts.append(para.text)
        for tbl_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_cells = [" ".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                             for cell in row.cells]
                table_data.append(row_cells)
            if not table_data: continue
            headers = table_data[0]
            clean = [{headers[i]:(row[i] if i<len(row) else "") for i in range(len(headers))}
                     for row in table_data[1:]]
            tables.append({"type":"table","table_index":tbl_idx,"headers":headers,
                           "rows":clean,"extraction_method":"docx_direct"})
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                try: embedded_images.append({"url":rel.target_ref,
                        "description":rel.target_ref.split("/")[-1]})
                except: pass
    except Exception as exc:
        log.error("DOCX extraction error %s: %s", docx_path.name, exc)
        return "", [], []
    return _fix_arabic("\n".join(text_parts)), embedded_images, tables


# ── Image OCR ─────────────────────────────────────────────────────────────────

def _preprocess_image(img):
    if img.mode != "L": img = img.convert("L")
    img = ImageEnhance.Contrast(img).enhance(2.0)
    img = img.filter(ImageFilter.SHARPEN)
    if img.width < 1000 or img.height < 1000:
        scale = max(1000/img.width, 1000/img.height)
        img = img.resize((int(img.width*scale), int(img.height*scale)),
                         Image.Resampling.LANCZOS)
    return img

def extract_image(image_path: Path) -> str:
    if not OCR_OK: return "[OCR unavailable – pytesseract/Pillow not installed]"
    try:
        img = _preprocess_image(Image.open(image_path))
        try: text = pytesseract.image_to_string(img, lang="eng+fra+ara")
        except: text = pytesseract.image_to_string(img, lang="eng")
        return _fix_arabic(text.strip()) if text.strip() else ""
    except Exception as exc:
        log.error("OCR error %s: %s", image_path.name, exc)
        return ""


# ── JSON builder ──────────────────────────────────────────────────────────────

def build_json(file_path: Path, archive_source: Optional[str] = None) -> Optional[Dict]:
    ext = file_path.suffix.lower()
    if ext not in EXTRACTABLE_EXT:
        return None

    if ext in PDF_EXT:    file_type = "pdf"
    elif ext in DOCX_EXT: file_type = "docx"
    elif ext in IMAGE_EXT: file_type = "image"
    elif ext in XLSX_EXT: file_type = "xlsx"
    else: return None

    log.info("  Extracting [%s] %s", file_type.upper(), file_path.name)
    text=""; pages=[]; tables=[]; images_list=[]

    if file_type == "pdf":
        text, pages, tables = extract_pdf(file_path)
    elif file_type == "docx":
        text, images_list, tables = extract_docx(file_path)
    elif file_type == "image":
        text = extract_image(file_path)
        tables = []
    elif file_type == "xlsx":
        text, tables = extract_xlsx(file_path)

    # ── Filter decision ───────────────────────────────────────────────────────
    keep, reason = should_keep(text, tables, file_path.name)
    if not keep:
        log.warning("  ✗ DROPPED %s — %s", file_path.name, reason)
        return None
    log.info("  ✓ KEPT %s — %s", file_path.name, reason)

    file_meta = {
        "path": str(file_path),
        "name": file_path.name,
        "type": file_type,
        "size_bytes": file_path.stat().st_size,
        "id": _file_id(file_path),
    }
    if archive_source:
        file_meta["archive_source"] = archive_source

    return {
        "metadata": {
            "file": file_meta,
            "extracted_at": datetime.now().isoformat(),
            "filter_reason": reason,
        },
        "content": {"text": text, "pages": pages, "tables": tables},
        "resources": {"links": [], "images": images_list, "documents": []},
    }


# ╔══════════════════════════════════════════════════════════════════════════╗
# ║  PIPELINE ORCHESTRATOR (SINGLE FACULTY)                                  ║
# ╚══════════════════════════════════════════════════════════════════════════╝

def run_pipeline_single(input_dir: Path, output_dir: Path) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)

    log.info("=" * 65)
    log.info("DOCUMENT EXTRACTION PIPELINE (Single Faculty)")
    log.info("Input  : %s", input_dir)
    log.info("Output : %s", output_dir)
    log.info("=" * 65)

    # ── Phase 0: Sort any files sitting loose in the root ─────────────────────
    log.info("\n[PHASE 0] Sorting root-level files into subfolders ...")
    _sort_root_files(input_dir)

    # ── Phase 1: Unpack archives (without deletion) ───────────────────────────
    log.info("\n[PHASE 1] Unpacking archives in %s/docs/ ...", input_dir.name)
    sorted_files = unpack_archives_keep(input_dir)
    log.info("  Phase 1 complete: %d file(s) copied into subfolders.", len(sorted_files))

    # ── Phase 2: Extract content from all organised files ─────────────────────
    log.info("\n[PHASE 2] Extracting & filtering content from all files ...")

    all_files: List[Path] = []
    for folder_name in ("pdfs", "docs", "images"):
        folder = input_dir / folder_name
        if folder.exists():
            for f in sorted(folder.rglob("*")):
                if f.is_file() and f.suffix.lower() in EXTRACTABLE_EXT:
                    all_files.append(f)

    # Legacy 'pages' folder support
    pages_folder = input_dir / "pages"
    if pages_folder.exists():
        for f in sorted(pages_folder.rglob("*")):
            if f.is_file() and f.suffix.lower() in EXTRACTABLE_EXT:
                all_files.append(f)

    if not all_files:
        log.warning("No extractable files found.")
        return

    log.info("  Found %d file(s) to process.", len(all_files))
    ok = dropped = fail = 0
    drop_reasons: Dict[str, int] = {}

    for file_path in all_files:
        try:
            data = build_json(file_path)
            if data is None:
                dropped += 1
                continue
            json_name = (
                file_path.stem + "_"
                + data["metadata"]["file"]["id"]
                + ".json"
            )
            out_path = output_dir / json_name
            with open(out_path, "w", encoding="utf-8") as fh:
                json.dump(data, fh, ensure_ascii=False, indent=2)
            log.info(
                "  ✓ → %s  (%d chars, %d table(s))",
                json_name,
                len(data["content"]["text"]),
                len(data["content"]["tables"]),
            )
            ok += 1
        except Exception as exc:
            log.error("  ✗ FAILED %s: %s", file_path.name, exc)
            fail += 1

    log.info("")
    log.info("=" * 65)
    log.info("Pipeline complete.")
    log.info("  Kept    : %d", ok)
    log.info("  Dropped : %d  (filtered out)", dropped)
    log.info("  Failed  : %d  (errors)", fail)
    log.info("Output : %s", output_dir)
    log.info("=" * 65)


# ╔══════════════════════════════════════════════════════════════════════════╗
# ║  ENTRY POINT                                                             ║
# ╚══════════════════════════════════════════════════════════════════════════╝

def main():
    parser = argparse.ArgumentParser(
        description="Extract content from documents of a single faculty (no deletion).")
    parser.add_argument("--input", "-i", default="./university_farhat_abaas")
    parser.add_argument("--output", "-o", default="extracted")
    parser.add_argument("--min-chars", type=int, default=MIN_TEXT_CHARS)
    args = parser.parse_args()

    root_dir = Path(args.input).resolve()
    if not root_dir.exists() or not root_dir.is_dir():
        print(f"Error: Input folder does not exist: {root_dir}")
        sys.exit(1)

    for subfolder in root_dir.iterdir():
        if subfolder.is_dir():
            print(f"\n🚀 Processing: {subfolder.name}")

            input_dir = subfolder
            output_dir = subfolder / args.output

            run_pipeline_single(input_dir, output_dir)


if __name__ == "__main__":
    main()